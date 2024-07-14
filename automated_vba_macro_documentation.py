import tkinter as tk
from tkinter import filedialog, messagebox
import os
import openpyxl
import re
from docx import Document
import spacy

# Load the English model for NLP
nlp = spacy.load("en_core_web_sm")

# Function to extract VBA code from an Excel file
def extract_vba_code(excel_file):
    vba_code = []
    try:
        wb = openpyxl.load_workbook(excel_file, keep_links=False)
        # Simulated extraction logic; replace with actual code extraction
        vba_code = [
            "Function CalculateSum(a As Integer, b As Integer) As Integer",
            "    Dim result As Integer",
            "    result = a + b",
            "    CalculateSum = result",
            "End Function",
            "Sub SampleMacro()",
            "    Dim x As Integer",
            "    x = 5",
            "    If x > 0 Then",
            "        MsgBox \"Positive Number\"",
            "    End If",
            "End Sub",
            "For i = 1 To 10",
            "    MsgBox i",
            "Next i"
        ]
    except Exception as e:
        raise RuntimeError(f"Failed to extract VBA code: {e}")
    return vba_code

# Function to extract components from VBA code
def extract_components(vba_code):
    functions = []
    subroutines = []
    variables = {}
    control_structures = []
    loops = []

    for line in vba_code:
        line = line.strip()

        if line.lower().startswith("function"):
            functions.append(line)
        elif line.lower().startswith("sub"):
            subroutines.append(line)

        var_match = re.findall(r'\bDim\s+(\w+)\s+As\s+(\w+)', line, re.IGNORECASE)
        for var in var_match:
            variables[var[0]] = var[1]

        if re.search(r'\b(If|For|While|Select|Do|With)\b', line, re.IGNORECASE):
            control_structures.append(line)
            if line.lower().startswith("for"):
                loops.append(line)
        elif "End If" in line or "End For" in line:
            control_structures.append(line)

    return functions, subroutines, variables, control_structures, loops

# Function to analyze VBA code using NLP
def analyze_vba_code(vba_code):
    descriptions = []
    
    for line in vba_code:
        doc = nlp(line)
        description = []

        for token in doc:
            if token.dep_ == "nsubj":
                description.append(f"The variable '{token.text}' is used.")
            elif token.dep_ == "ROOT":
                description.append(f"This line performs the action: '{token.text}'.")

        descriptions.append(" ".join(description) or "No specific action identified.")
    
    return descriptions

# Function to generate the Word document
def generate_word_document(functions, subroutines, variables, control_structures, loops, vba_code, vba_descriptions):
    doc = Document()
    doc.add_heading('VBA Macro Documentation', level=1)

    doc.add_heading('Functions:', level=2)
    for f in functions:
        doc.add_paragraph(f"- **{f}**\n  - **Returns:** Integer\n  - **Parameters:** {f.split('(')[1].split(')')[0]}\n  - **Example Usage:** result = {f.split(' ')[1]}(1, 2)")

    doc.add_heading('Subroutines:', level=2)
    for s in subroutines:
        doc.add_paragraph(f"- **{s}**\n  - **Parameters:** None\n  - **Example Usage:** Call {s.split(' ')[1]}")

    doc.add_heading('Variables:', level=2)
    for v, vt in variables.items():
        doc.add_paragraph(f"- **{v}** (Declared as {vt})")

    doc.add_heading('Control Structures:', level=2)
    for cs in control_structures:
        doc.add_paragraph(f"- {cs}")

    doc.add_heading('Loops:', level=2)
    for l in loops:
        doc.add_paragraph(f"- {l}")

    doc.add_heading('VBA Code Analysis:', level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'VBA Code'
    hdr_cells[1].text = 'Description'

    for line, desc in zip(vba_code, vba_descriptions):
        row_cells = table.add_row().cells
        row_cells[0].text = line
        row_cells[1].text = desc

    save_path = os.path.join(os.getcwd(), "vba_documentation.docx")
    doc.save(save_path)

    os.startfile(save_path)

# Function to run the documentation generation
def run_documentation():
    excel_file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsm")])
    if not excel_file:
        return

    try:
        vba_code = extract_vba_code(excel_file)
        functions, subroutines, variables, control_structures, loops = extract_components(vba_code)

        vba_descriptions = analyze_vba_code(vba_code)

        generate_word_document(functions, subroutines, variables, control_structures, loops, vba_code, vba_descriptions)

        messagebox.showinfo("Success", "Documentation generated: vba_documentation.docx")
    except RuntimeError as e:
        messagebox.showerror("Error", str(e))
    except Exception as e:
        messagebox.showerror("Error", f"An unexpected error occurred: {e}")

# Function to close the application
def close_application():
    root.destroy()

# Function to set up the GUI
def setup_gui():
    global root
    root = tk.Tk()
    root.title("VBA Macro Documentation Tool")
    root.attributes('-fullscreen', True)

    # Gradient background
    root.configure(bg="#2e1a47")

    # Canvas for gradient effect
    canvas = tk.Canvas(root, bg="#2e1a47", height=800, width=1200)
    canvas.pack(fill=tk.BOTH, expand=True)

    # Create gradient background
    for i in range(1200):
        color = f'#{int(46 + (105 - 46) * (i / 1200)):02x}{int(26 + (10 - 26) * (i / 1200)):02x}{int(71 + (206 - 71) * (i / 1200)):02x}'
        canvas.create_line(i, 0, i, 800, fill=color)

    # Description label
    description = tk.Label(root, 
                            text=("This tool allows users to generate comprehensive documentation for VBA macros present in Excel files. "
                                  "Simply select a macro-enabled Excel file, and the tool will analyze the VBA code to produce a well-structured "
                                  "Word document detailing functions, subroutines, variables, control structures, loops, and a line-by-line analysis."),
                            bg="#2e1a47", 
                            fg="#ffccff", 
                            font=("Arial", 16), 
                            wraplength=800)
    description.place(relx=0.5, rely=0.3, anchor="center")

    # Generate button
    btn_generate = tk.Button(root, text="Generate Documentation", command=run_documentation, 
                             bg="#ffccff", fg="black", font=("Arial", 14, "bold"), relief="raised", bd=5)
    btn_generate.place(relx=0.5, rely=0.5, anchor="center")

    # Close button
    btn_close = tk.Button(root, text="Back to dashboard", command=close_application, 
                          bg="#ffccff", fg="black", font=("Arial", 14, "bold"), relief="raised", bd=5)
    btn_close.place(relx=0.5, rely=0.65, anchor="center")

    # Button hover effect
    btn_generate.bind("<Enter>", lambda e: btn_generate.config(bg="#ff99ff"))
    btn_generate.bind("<Leave>", lambda e: btn_generate.config(bg="#ffccff"))

    btn_close.bind("<Enter>", lambda e: btn_close.config(bg="#ff99ff"))
    btn_close.bind("<Leave>", lambda e: btn_close.config(bg="#ffccff"))

    root.mainloop()

if __name__ == "__main__":
    setup_gui()
